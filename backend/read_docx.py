# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

path = "D:/桌面/rg课设/软件工程课程设计任务书_2026春/选题指南/T-01_NekoCafé智慧餐饮预约平台.docx"
doc = Document(path)

print("=" * 60)
print("【段落内容】")
print("=" * 60)
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"[{i}] {para.text}")

print()
print("=" * 60)
print("【表格内容】")
print("=" * 60)
for t_idx, table in enumerate(doc.tables):
    print(f"\n--- 表格 {t_idx+1} ---")
    for row in table.rows:
        row_data = [cell.text.strip() for cell in row.cells]
        print(" | ".join(row_data))
